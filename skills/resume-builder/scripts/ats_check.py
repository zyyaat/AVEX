#!/usr/bin/env python3
"""
ats_check.py — 简历 ATS（Applicant Tracking System）友好度检查 + 关键词覆盖率

用法：
    python ats_check.py --resume resume.md --industry internet
    python ats_check.py --resume resume.md --industry tech --jd jd.txt
    python ats_check.py --resume resume.docx --industry finance --out report.md

支持输入：.md / .txt / .docx（docx 走 python-docx，需要 pip install python-docx）
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

INDUSTRIES = {"internet", "tech", "finance", "general"}


def load_resume_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            print(
                "✗ 缺少 python-docx，请先安装：pip install python-docx --break-system-packages",
                file=sys.stderr,
            )
            sys.exit(1)
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    print(f"✗ 暂不支持的格式：{suffix}", file=sys.stderr)
    sys.exit(1)


def load_keywords(industry: str, references_dir: Path) -> list[str]:
    path = references_dir / "keywords" / f"{industry}.txt"
    if not path.exists():
        print(f"✗ 关键词库不存在：{path}", file=sys.stderr)
        sys.exit(1)
    return [
        line.strip()
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]


def coverage(resume_text: str, keywords: list[str]) -> dict:
    text_lower = resume_text.lower()
    hits, missing = [], []
    for kw in keywords:
        # 兼容大小写、中英混排
        if kw.lower() in text_lower:
            hits.append(kw)
        else:
            missing.append(kw)
    return {
        "hits": hits,
        "missing": missing,
        "rate": len(hits) / len(keywords) if keywords else 0,
    }


def jd_extract_keywords(jd_text: str) -> list[str]:
    """从 JD 文本里抽取候选关键词。简易版：取常见技能/工具/动词。"""
    # 抓中文 2~10 字、英文 2~30 字的"实词"
    candidates = re.findall(
        r"[A-Za-z][A-Za-z0-9+/.\-_]{1,29}|[一-龥]{2,10}",
        jd_text,
    )
    # 简单去停用词
    stop = {
        "公司", "工作", "我们", "你将", "团队", "需要", "能够", "具备", "熟悉",
        "了解", "良好", "优秀", "经验", "能力", "岗位", "职责", "要求", "以上",
        "相关", "及其", "或者", "进行", "完成", "负责", "推动", "实现", "提升",
        "并且", "包括", "以下", "根据",
    }
    seen = set()
    out = []
    for c in candidates:
        key = c.lower()
        if key in seen or c in stop:
            continue
        seen.add(key)
        out.append(c)
    return out[:80]  # 取前 80 个候选


def ats_friendliness(text: str, source_path: Path) -> tuple[int, list[str]]:
    """评估 ATS 友好度，返回 (分数 / 10, 警告列表)。"""
    score = 10
    warnings = []

    # 长度
    if len(text) < 200:
        score -= 3
        warnings.append("⚠ 简历文本过短（< 200 字符），可能解析失败或内容不足")
    elif len(text) > 6000:
        score -= 1
        warnings.append("⚠ 简历偏长（> 6000 字符），建议精简到 1~2 页")

    # 联系方式
    has_email = bool(re.search(r"[\w.\-+]+@[\w.\-]+\.\w+", text))
    has_phone = bool(re.search(r"(\+?86[-\s]?)?1[3-9]\d{9}|\d{3}[-\s]?\d{4}[-\s]?\d{4}", text))
    if not has_email:
        score -= 1
        warnings.append("⚠ 没找到邮箱")
    if not has_phone:
        score -= 1
        warnings.append("⚠ 没找到手机号")

    # 装饰符号
    decorative = re.findall(r"[★☆●○◆◇▶▷■□▪▫♦]", text)
    if len(decorative) > 5:
        score -= 1
        warnings.append(f"⚠ 用了 {len(decorative)} 个装饰符号（★●◆等），ATS 可能识别异常，建议精简")

    # emoji
    emojis = re.findall(r"[\U0001F300-\U0001FAFF\U0001F600-\U0001F64F]", text)
    if emojis:
        score -= 1
        warnings.append(f"⚠ 检测到 {len(emojis)} 个 emoji，部分 ATS 会乱码，建议删掉")

    # 数字密度（量化结果是否充分）
    numbers = re.findall(r"\d+(?:\.\d+)?%?", text)
    bullet_count = len(re.findall(r"^\s*[-*•]\s+", text, flags=re.MULTILINE))
    if bullet_count > 0:
        density = len(numbers) / bullet_count
        if density < 0.4:
            score -= 1
            warnings.append(
                f"⚠ 量化密度低（每条 bullet 平均 {density:.2f} 个数字），"
                f"建议在工作 / 项目经历里多加数字"
            )

    # docx 特定
    if source_path.suffix.lower() == ".docx":
        try:
            from docx import Document
            doc = Document(str(source_path))
            tables = len(doc.tables)
            images = sum(1 for s in doc.inline_shapes)
            if tables > 1:
                score -= 1
                warnings.append(
                    f"⚠ docx 里有 {tables} 个表格，部分 ATS 解析表格会丢字段，"
                    f"建议改成正文段落"
                )
            if images > 0:
                score -= 1
                warnings.append(
                    f"⚠ docx 里有 {images} 张图片（含证件照），ATS 不读图，"
                    f"重要信息别只放在图里；证件照可保留"
                )
        except Exception:
            pass

    return max(0, score), warnings


def render_report(
    industry_cov: dict,
    industry: str,
    jd_cov: dict | None,
    ats_score: int,
    ats_warnings: list[str],
) -> str:
    lines = [
        "# 简历 ATS 检查报告",
        "",
        f"## 行业关键词覆盖（{industry}）",
        f"- 命中率：**{industry_cov['rate'] * 100:.1f}%** "
        f"({len(industry_cov['hits'])} / {len(industry_cov['hits']) + len(industry_cov['missing'])})",
        "",
        "**已命中**：" + (", ".join(industry_cov["hits"]) or "（无）"),
        "",
        "**建议补充**（前 15 个）：" + (", ".join(industry_cov["missing"][:15]) or "（无）"),
        "",
    ]

    if jd_cov is not None:
        lines += [
            "## JD 关键词覆盖",
            f"- 命中率：**{jd_cov['rate'] * 100:.1f}%**",
            "",
            "**已命中**：" + (", ".join(jd_cov["hits"][:30]) or "（无）"),
            "",
            "**JD 出现但简历没有**（重点补这些）：" + (", ".join(jd_cov["missing"][:20]) or "（无）"),
            "",
        ]

    lines += [
        f"## ATS 友好度评分：**{ats_score}/10**",
        "",
    ]
    if ats_warnings:
        lines += ats_warnings
    else:
        lines.append("✅ 没有明显问题")

    lines += [
        "",
        "---",
        "## 改进建议优先级",
        "",
        "1. 先补 JD 命中率 → 这是 ATS 通过率的最直接信号",
        "2. 再补行业关键词 → 让简历能在更宽的搜索里被捞到",
        "3. 最后调 ATS 友好度 → 移除装饰符号、emoji、表格、图片",
        "",
        "注意：覆盖率不是越高越好，**关键词必须出现在真实的成就 bullet 里**，",
        "不要把关键词单独列一长串当 skills，会被 HR 一眼识破。",
    ]
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--resume", required=True, help="简历文件路径 (.md/.txt/.docx)")
    parser.add_argument(
        "--industry", choices=list(INDUSTRIES), required=True, help="行业方向"
    )
    parser.add_argument("--jd", help="可选：JD 文本文件，做精准对比")
    parser.add_argument("--out", help="输出报告路径，缺省直接打印")
    parser.add_argument(
        "--references-dir",
        default=str(Path(__file__).resolve().parent.parent / "references"),
    )
    args = parser.parse_args()

    resume_path = Path(args.resume).expanduser()
    if not resume_path.exists():
        print(f"✗ 简历文件不存在：{resume_path}", file=sys.stderr)
        sys.exit(1)

    text = load_resume_text(resume_path)
    references_dir = Path(args.references_dir)

    industry_keywords = load_keywords(args.industry, references_dir)
    industry_cov = coverage(text, industry_keywords)

    jd_cov = None
    if args.jd:
        jd_path = Path(args.jd).expanduser()
        if not jd_path.exists():
            print(f"✗ JD 文件不存在：{jd_path}", file=sys.stderr)
            sys.exit(1)
        jd_keywords = jd_extract_keywords(jd_path.read_text(encoding="utf-8"))
        jd_cov = coverage(text, jd_keywords)

    ats_score, ats_warnings = ats_friendliness(text, resume_path)
    report = render_report(industry_cov, args.industry, jd_cov, ats_score, ats_warnings)

    if args.out:
        out_path = Path(args.out).expanduser()
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(report, encoding="utf-8")
        print(f"✓ 报告已生成：{out_path}")
    else:
        print(report)


if __name__ == "__main__":
    main()
